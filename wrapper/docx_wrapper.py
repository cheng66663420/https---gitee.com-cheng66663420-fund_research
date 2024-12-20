import os

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxWrapper:
    def __init__(self, file_name: str = None):
        self.document = Document()
        self.file_name = file_name

    def __hex_color2_rgb(self, hex_color: str) -> tuple:
        hex_color = hex_color.strip("#")
        return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))

    def save_document(self):
        """
        保存文档
        """
        self.document.save(self.file_name)

    def _set_paragraph_font(
        self,
        paragraph: Paragraph,
        text: str,
        **kwargs,
    ):
        """
        设置段落字体

        Parameters
        ----------
        paragraph : Paragraph
            段落
        text : str
            文本
        ----------
        kwargs : dict
        ----------
            字体属性
            name : str
                字体名称, 默认楷体
            size : int
                字体大小, 默认18

            color : str
                字体颜色，格式为#RRGGBB,默认#000000
            bold : bool
                是否加粗, 默认False
            italic : bool
                是否斜体, 默认False
            before_space : float
                段前间距，单位为磅, 默认0
            after_space : float
                段后间距，单位为磅, 默认0
            line_space : float
                行间距，单位为行距, 默认1倍行距
            first_line_indent : float
                首行缩进，单位为字符, 默认0
            alignment : int
                对齐方式, 0左对齐, 1居中, 2右对齐, 3两端 默认0

        """
        default_font_dict = {
            "name": "楷体",
            "size": 18,
            "color": "#000000",
            "bold": False,
            "italic": False,
            "before_space": 0,
            "after_space": 0,
            "line_space": 1,
            "first_line_indent": 0,
            "alignment": 1,
        }
        for key in kwargs:
            if key not in default_font_dict:
                raise ValueError(f"Invalid font key: {key}")
        default_font_dict |= kwargs
        paragraph.paragraph_format.space_before = Pt(default_font_dict["before_space"])
        paragraph.paragraph_format.space_after = Pt(default_font_dict["after_space"])
        paragraph.paragraph_format.line_spacing = default_font_dict["line_space"]
        paragraph.paragraph_format.first_line_indent = Pt(
            default_font_dict["first_line_indent"]
        )
        paragraph.alignment = default_font_dict["alignment"]
        run = paragraph.add_run(text)
        run.font.name = default_font_dict["name"]
        run.font.size = Pt(default_font_dict["size"])
        r, g, b = self.__hex_color2_rgb(default_font_dict["color"])
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.bold = default_font_dict["bold"]
        run.font.italic = default_font_dict["italic"]
        run._element.rPr.rFonts.set(qn("w:eastAsia"), default_font_dict["name"])

    def add_paragraph(self, text: str, **kwargs):
        """
        增加一个段落,并设置格式
        Parameters
        ----------
        text : str
            段落文本
        ----------
        kwargs : dict
        ----------
            字体属性
            name : str
                字体名称, 默认楷体
            size : int
                字体大小, 默认18
            color : str
                字体颜色，格式为#RRGGBB,默认#000000
            bold : bool
                是否加粗, 默认False
            italic : bool
                是否斜体, 默认False
            before_space : float
                段前间距，单位为磅, 默认0
            after_space : float
                段后间距，单位为磅, 默认0
            line_space : float
                行间距，单位为行距, 默认1倍行距
            first_line_indent : float
                首行缩进，单位为字符, 默认0
            alignment : int
                对齐方式, 0左对齐, 1居中, 2右对齐, 3两端 默认0
        """
        p = self.document.add_paragraph()
        self._set_paragraph_font(p, text, **kwargs)

    def add_heading(self, text: str, level: int = 1, **kwargs):
        """
        增加标题

        Parameters
        ----------
        text : str
            标题文本
        level : int, optional
            标题级别, by default 1
        ----------
        kwargs : dict
        ----------
            字体属性
            name : str
                字体名称, 默认楷体
            size : int
                字体大小, 默认18
            color : str
                字体颜色，格式为#RRGGBB,默认#000000
            bold : bool
                是否加粗, 默认False
            italic : bool
                是否斜体, 默认False
            before_space : float
                段前间距，单位为磅, 默认0
            after_space : float
                段后间距，单位为磅, 默认0
            line_space : float
                行间距，单位为行距, 默认1倍行距
            first_line_indent : float
                首行缩进，单位为字符, 默认0
            alignment : int
                对齐方式, 0左对齐, 1居中, 2右对齐, 3两端 默认0
        """
        head = self.document.add_heading(level=level)
        self._set_paragraph_font(head, text, **kwargs)

    def add_table(self, data: list, style: str = None) -> Table:
        """
        讲数据写入表格

        Parameters
        ----------
        data : list
            数据二维列表
        rows : int
            行数
        cols : int
            _description_
        style : str, optional
            _description_, by default None

        Returns
        -------
        Table
            _description_
        """
        data_array = np.array(data)
        rows, cols = data_array.shape
        table = self.document.add_table(rows=rows, cols=cols, style=style)
        for i in range(rows):
            for j in range(cols):
                table.cell(i, j).text = str(data_array[i][j])

        self.set_table_boarder(table)
        self._set_table_style(table)
        return table

    def add_image(self, image_path, width=Inches(1.25)):
        """
        添加图片
        :param image_path: 图片路径
        :param width: 图片宽度，默认为1.25英寸
        """
        if not os.path.isfile(image_path):
            print(f"图片文件不存在: {image_path}")
            return
        self.document.add_picture(image_path, width=width)

    def add_dataframe(self, df: pd.DataFrame, header: bool = True, index: bool = False):
        """
        将DataFrame写入Word文档
        :param df: DataFrame对象
        :param header: 是否包含表头，默认True
        :param index: 是否包含索引，默认False
        """
        df = df.copy()
        if index:
            df = df.reset_index()
        if header:
            df1 = pd.DataFrame(df.columns.tolist()).T
            df1.columns = df.columns
            df = pd.concat([df1, df]).reset_index(drop=True)
        data_array = df.values
        return self.add_table(data_array)

    def _set_table_style(self, table: Table):

        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                cell = table.cell(row, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"  # 设置英文字体
                cell.paragraphs[0].runs[0].font.size = Pt(12)  # 字体大小
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(
                    qn("w:eastAsia"), "楷体"
                )
                cell.paragraphs[0].paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )

    def _set_table_single_boarder(self, table: Table, **kwargs):
        """
        Set table`s border
        Usage:
        set_table_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            left={"sz": 24, "val": "dashed"},
            right={"sz": 12, "val": "dashed"},
        )
        """
        borders = OxmlElement("w:tblBorders")
        for tag in ("bottom", "top", "left", "right", "insideV", "insideH"):
            if edge_data := kwargs.get(tag):
                any_border = OxmlElement(f"w:{tag}")
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        any_border.set(qn(f"w:{key}"), str(edge_data[key]))
                borders.append(any_border)
                table._tbl.tblPr.append(borders)

    def set_table_boarder(self, table: Table):
        """为表格添加边框"""
        return self._set_table_single_boarder(
            table,
            top={"sz": 4, "val": "single", "color": "#000000"},
            bottom={"sz": 4, "val": "single", "color": "#000000"},
            left={"sz": 4, "val": "single", "color": "#000000"},
            right={"sz": 4, "val": "single", "color": "#000000"},
            insideV={"sz": 4, "val": "single", "color": "#000000"},
            insideH={"sz": 4, "val": "single", "color": "#000000"},
        )

    def merge_table_cells(
        self,
        table: Table,
        row_index: int,
        col_index: int,
        row2merge: int,
        col2merge: int,
        text: str = None,
    ):
        cell1 = table.cell(row_index, col_index)
        if text is None:
            text = cell1.text
        cell2 = table.cell(row2merge, col2merge)
        cell1.merge(cell2)
        cell1.text = text


# 使用示例
if __name__ == "__main__":
    doc_gen = DocxWrapper("e:/example.docx")

    # 添加段落
    doc_gen.add_paragraph("这是一个示例段落。", bold=1, alignment=0)

    # 添加表格
    data = [["Name", "Age"], ["Alice", 24], ["Bob", 30]]
    doc_gen.add_table(data=data)

    # # 添加图片（请确保图片路径正确）
    # doc_gen.add_image("path_to_your_image.jpg")
    # 添加标题
    doc_gen.add_heading("示例标题", level=1, bold=1, size=16, alignment=1)
    # 将DataFrame转换为表格
    df = pd.DataFrame({"Column1": [1, 2, 3], "Column2": ["A", "B", "C"]})
    doc_gen.add_dataframe(df, header=True, index=False)

    # 保存文档
    doc_gen.save_document()
