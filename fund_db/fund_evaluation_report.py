import datetime
import math
import os
import sys
import warnings

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from quant_utils.df2wordtable import DF2WordTable
from quant_utils.utils import make_dirs

warnings.simplefilter(action="ignore", category=FutureWarning)
from fund_db.fund_evation_fuction import *


class FundEvaluation:
    def __init__(
        self,
        ticker_symbol,
        end_date,
        include_out_tranche=0,
        file_path="E:/持续评估/",
        if_new=0,
        tranche="优选池",
    ):
        self.ticker_symbol = ticker_symbol
        self.end_date = end_date
        self.file_path = file_path

        self.if_new = if_new
        self.include_out_tranche = include_out_tranche
        self.tranche = tranche
        if if_new == 1:
            self.include_out_tranche = 1

        make_dirs(self.file_path)
        self.fund_type = get_fund_type(ticker_symbol=ticker_symbol, end_date=end_date)
        self.fund_type_dict = {
            "LEVEL_1": self.fund_type["LEVEL_1"].values[0],
            "LEVEL_2": self.fund_type["LEVEL_2"].values[0],
            "LEVEL_3": self.fund_type["LEVEL_3"].values[0],
            "LEVEL_SUM": self.fund_type["LEVEL_SUM"].values[0],
        }

        self.template_name = (
            "指数和商品"
            if (self.fund_type_dict["LEVEL_SUM"] in ("商品", "指数"))
            or (
                self.fund_type_dict["LEVEL_2"]
                in ("指数增强", "国际(QDII)增强指数型股票基金")
            )
            else self.fund_type_dict["LEVEL_SUM"]
        )

        self.type = (
            "指数增强"
            if (
                self.fund_type_dict["LEVEL_2"]
                in ("指数增强", "国际(QDII)增强指数型股票基金")
            )
            else self.fund_type_dict["LEVEL_SUM"]
        )
        self.related_code = (
            ",".join(
                get_related_code_list(self.ticker_symbol, self.include_out_tranche)
            )
            if self.type != "货币"
            else self.ticker_symbol
        )
        self.start_date = (
            get_evaluated_date(self.ticker_symbol, self.end_date)
            if self.template_name in ("主动权益", "主动债券")
            else get_other_evaluated_date(self.ticker_symbol, self.end_date)
        )
        self.fund_df = pd.read_excel(
            f"{file_path}/基金评分模板.xlsx",
            sheet_name=self.template_name,
            index_col=[0, 1, 2],
        )

        name_dict = get_fund_name_dict(ticker_symbol)
        self.fund_name = name_dict["简称"] if self.type == "货币" else name_dict["全称"]

        # 基金公司
        self.management_company = get_fund_management_company(ticker_symbol)

        # 基金公司策略
        self.fund_company_type_df = get_fund_type_comapany_rank(
            self.ticker_symbol, self.end_date
        )

        # 基金公司策略排名
        self.fund_company_type_rank = self.fund_company_type_df["排名"].values[0]

        # 基金公司策略规模
        self.fund_company_type_asset = self.fund_company_type_df["NET_ASSET"].values[0]

        # 基金公司近一年离职率
        self.fund_company_stablility = get_fund_company_stablility(
            self.ticker_symbol, self.end_date
        )

        # 基金经理
        self.fund_manager_df = get_fund_manager(self.ticker_symbol, self.end_date)
        manager_duration = self.fund_manager_df["MANAGER_DURATION"].values[0].split(";")
        self.manager_duration = max([float(i) for i in manager_duration])
        self.manager_name = self.fund_manager_df["MANAGER_NAME"].values[0]

        # 基金规模
        # 如果是etf联接则取etf的规模
        self.fund_asset = (
            get_fund_asset(self.ticker_symbol, self.end_date)
            if self.fund_type_dict["LEVEL_2"] != "ETF联接"
            else get_fund_asset(get_etf_fund_code(self.ticker_symbol), self.end_date)
        )

        self.fund_annualised_alpha_dict = get_fund_annualised_alpha(
            self.ticker_symbol, self.start_date, self.end_date
        )
        self.fund_perf_rank_dict = get_fund_perf_rank(self.ticker_symbol, self.end_date)
        if self.if_new == 1:
            self.fund_in_tranch_dates = 0
        else:
            self.fund_in_tranch_dates = float(
                get_fund_in_tranch_dates(self.ticker_symbol)["入池时间"].values[0]
            )

        # 基金公司部分
        self.fund_company_desc = get_fund_company_desc(self.management_company)

        # 基金成立年限
        self.fund_establish_duration = get_fund_establish_duration(
            self.ticker_symbol, self.end_date
        )
        self.fund_manager_asset = get_fund_manager_fund_asset(
            self.ticker_symbol, self.end_date
        )

        # 存储路径
        self.save_path = f"{self.file_path}/{self.end_date}/{self.template_name}-{self.fund_name.replace('/', '')}-{self.manager_name}"
        self.evaluation_result = "满足"
        self.remark = []

    def _add_paragraph(self, document: Document, text: str):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(text)
        run.font.name = "楷体"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

    def _merge_cells(
        self, table, row_index, col_index, row2merge, col2merge, text=None
    ):
        cell1 = table.cell(row_index, col_index)
        if text is None:
            text = cell1.text
        cell2 = table.cell(row2merge, col2merge)
        cell1.merge(cell2)
        cell1.text = text

    def _set_table_style(self, table):
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                cell = table.cell(row, col)
                cell.paragraphs[0].runs[0].font.name = "Times New Roman"  # 设置英文字体
                cell.paragraphs[0].runs[0].font.size = Pt(12)  # 字体大小
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(
                    qn("w:eastAsia"), "楷体"
                )  # 设置中文字体

    def info_part(self):
        fund_name = (
            get_related_code_str(self.ticker_symbol, self.include_out_tranche)
            if self.type != "货币"
            else f"{self.ticker_symbol}:{self.fund_name}"
        )
        self.fund_df.loc[("信息", "基金产品", None), "简要评价"] = fund_name
        self.fund_df.loc[("信息", "基金经理", None), "简要评价"] = self.manager_name
        self.fund_df.loc[("信息", "评估日期", None), "简要评价"] = (
            f"{self.start_date}-{self.end_date}"
        )

    def company_part(self):
        # 基金公司部分
        fund_company_str = (
            f"""{self.management_company}{self.type}基金产品管理规模"""
            + f"""{self.fund_company_type_asset:.2f}亿元,"""
            + f"""在基金公司中排名{self.fund_company_type_rank}"""
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "简要评价"] = (
            fund_company_str
        )

        company_rank_score = (
            10
            if self.fund_company_type_rank <= 10
            else (
                8
                if self.fund_company_type_rank <= 20
                else (5 if self.fund_company_type_rank <= 50 else 0)
            )
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "分数"] = (
            company_rank_score
        )

    def manager_part(self):
        # 基金经理部分
        manager_duration_str = f"""基金经理投资年限{self.manager_duration:.2f}年"""
        self.fund_df.loc[("定量分析", "基金经理", "基金经理投资年限"), "简要评价"] = (
            manager_duration_str
        )
        manager_duration_score = (
            0
            if self.manager_duration <= 1
            else (5 if self.manager_duration > 5 else self.manager_duration)
        )
        self.fund_df.loc[("定量分析", "基金经理", "基金经理投资年限"), "分数"] = (
            manager_duration_score
        )

    def qualitative_part(self):
        self.fund_df.loc[("定性分析", "基金公司", "投研团队建设"), "分数"] = 5
        self.fund_df.loc[("定性分析", "基金公司", "投研人员激励措施"), "分数"] = 5
        self.fund_df.loc[("定性分析", "基金公司", "合规表现情况"), "分数"] = 10
        self.fund_df.loc[("定性分析", "基金公司", "基金经理稳定性"), "分数"] = (
            self.fund_company_stablility["得分"]
        )
        self.fund_df.loc[("定性分析", "基金公司", "基金经理稳定性"), "简要评价"] = (
            f'公司近一年离职率{self.fund_company_stablility["近一年离职率"]*100:.2f}%'
        )
        self.fund_df.loc[
            ("定性分析", "基金经理", "投资逻辑框架是否有效；投资风格及稳定性"), "分数"
        ] = 15
        self.fund_df.loc[
            ("定性分析", "基金经理", "投资逻辑框架是否有效；投资风格及稳定性"),
            "简要评价",
        ] = "基金经理投资逻辑清晰,风格稳定,严格遵守投资约束,不存在漂移现象"
        if self.fund_company_desc.empty:
            self.fund_df.loc[("定性分析", "基金公司", "投研团队建设"), "简要评价"] = (
                "已经搭建完备的投资研究体系，核心投研团队负责人投资经验丰富，历经多轮市场牛熊"
            )
            self.fund_df.loc[
                ("定性分析", "基金公司", "投研人员激励措施"), "简要评价"
            ] = "已建立有效的员工激励制度，定性与定量考核激励指标相结合，注重中长期业绩考核评价。"
            self.fund_df.loc[("定性分析", "基金公司", "合规表现情况"), "简要评价"] = (
                "公司近三年合规表现良好，具备较为完善的内控机制"
            )
        # 定性分析部分
        else:
            self.fund_df.loc[("定性分析", "基金公司", "投研团队建设"), "简要评价"] = (
                self.fund_company_desc["TEAM_BUILD"].values[0]
            )
            self.fund_df.loc[
                ("定性分析", "基金公司", "投研人员激励措施"), "简要评价"
            ] = self.fund_company_desc["ENGAGEMENT"].values[0]
            self.fund_df.loc[("定性分析", "基金公司", "合规表现情况"), "简要评价"] = (
                self.fund_company_desc["COMPLIANCE"].values[0]
            )

        self.fund_df.loc[("加分项", "基金产品", "产品入库时间加分"), "分数"] = (
            0
            if self.fund_in_tranch_dates < 1
            else (10 if self.fund_in_tranch_dates > 2 else 5)
        )
        self.fund_df.loc[("加分项", "基金产品", "产品入库时间加分"), "简要评价"] = (
            f"产品入库时间{self.fund_in_tranch_dates:.2f}年"
        )

    def other_part(self):

        self.fund_df.loc[("其他", "其他", "基金成立时间是否满1年"), "分数"] = 0
        self.fund_df.loc[
            ("其他", "其他", "最新报告期基金净资产是否超过2亿"), "分数"
        ] = 0

        if (self.fund_establish_duration < 1) and (
            self.fund_type_dict["LEVEL_SUM"] != "指数"
        ):
            self.evaluation_result = "不满足"
            self.remark.append("非指数产品不满足成立1年")

        if self.fund_asset < 2:
            self.evaluation_result = "不满足"
            self.remark.append("基金资产不满2亿")

        self.fund_df.loc[("其他", "其他", "基金成立时间是否满1年"), "简要评价"] = (
            "基金产品不满一年"
            if (self.fund_establish_duration < 1)
            and (self.fund_type_dict["LEVEL_SUM"] != "指数")
            else f"基金产品成立{self.fund_establish_duration:.2f},大于1年"
        )

        self.fund_df.loc[
            ("其他", "其他", "最新报告期基金净资产是否超过2亿"), "简要评价"
        ] = (
            "最新报告期基金净资产不足2亿"
            if self.fund_asset < 2
            else f"最新报告期基金净资产{self.fund_asset:.2f}亿,大于2亿"
        )

        if self.template_name in ("主动权益", "主动债券"):
            self.fund_df.loc[("其他", "其他", "同策略产品管理规模"), "分数"] = 0
            self.fund_df.loc[("其他", "其他", "基金经理投资年限是否满1年"), "分数"] = 0
            if self.template_name == "主动权益":
                if self.tranche == "优选池":
                    if self.fund_manager_asset < 5:
                        self.evaluation_result = "不满足"
                        self.remark.append("优选池同策略产品管理规模不足5亿")

                    self.fund_df.loc[
                        ("其他", "其他", "同策略产品管理规模"), "简要评价"
                    ] = (
                        "同策略产品管理规模不足5亿"
                        if self.fund_manager_asset < 5
                        else f"同策略产品管理规模{self.fund_manager_asset:.2f}亿,大于5亿"
                    )
                else:
                    if self.fund_manager_asset < 20:
                        self.evaluation_result = "不满足"
                        self.remark.append("核心池同策略产品管理规模不足20亿")
                    self.fund_df.loc[
                        ("其他", "其他", "同策略产品管理规模"), "简要评价"
                    ] = (
                        "核心池同策略产品管理规模不足20亿"
                        if self.fund_manager_asset < 20
                        else f"核心池同策略产品管理规模{self.fund_manager_asset:.2f}亿,大于20亿"
                    )

            else:
                if self.tranche == "优选池":
                    if self.fund_manager_asset < 20:
                        self.evaluation_result = "不满足"
                        self.remark.append("优选池同策略产品管理规模不足20亿")
                    self.fund_df.loc[
                        ("其他", "其他", "同策略产品管理规模"), "简要评价"
                    ] = (
                        "优选池同策略产品管理规模不足20亿"
                        if self.fund_manager_asset < 20
                        else f"优选池同策略产品管理规模{self.fund_manager_asset:.2f}亿,大于20亿"
                    )
                else:
                    if self.fund_manager_asset < 50:
                        self.evaluation_result = "不满足"
                        self.remark.append("核心池同策略产品管理规模不足50亿")
                    self.fund_df.loc[
                        ("其他", "其他", "同策略产品管理规模"), "简要评价"
                    ] = (
                        "核心池同策略产品管理规模不足50亿"
                        if self.fund_manager_asset < 50
                        else f"优选池同策略产品管理规模{self.fund_manager_asset:.2f}亿,大于50亿"
                    )
            # self.fund_df.loc[("其他", "其他", "基金经理投资年限是否满1年"), "简要评价"] = (
            #     '基金经理投资年限不满一年' if self.manager_duration < 1
            #     else f'基金经理年限{self.manager_duration:.2f},大于1年'
            # )

    def save_excel(self):
        file_path = f"{self.file_path}/{self.end_date}"
        make_dirs(file_path)
        with pd.ExcelWriter(f"{self.save_path}.xlsx", engine="xlsxwriter") as writer:
            # workbook = writer.book
            # highlight_fmt = workbook.add_format({'bg_color': '#FFD7E2', 'num_format': '0.00'})
            # l_end = len(df.index) + 2 # 表格的行数,便于下面设置格式
            self.fund_df.to_excel(writer, "Sheet1")  # startcol=0, startrow=2
            workbook1 = writer.book
            worksheets = writer.sheets
            worksheet1 = worksheets["Sheet1"]
            # format1 = workbook1.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006'})
            format2 = workbook1.add_format(
                {
                    "bold": True,
                    "align": "center",
                    "valign": "top",
                    "fg_color": "white",
                    "text_wrap": False,
                }
            )

            bold1 = workbook1.add_format(
                {
                    "bold": True,  # 字体加粗
                    "border": 1,  # 单元格边框宽度
                    "align": "center",  # 水平对齐方式
                    "valign": "vcenter",  # 垂直对齐方式
                    "text_wrap": True,  # 是否自动换行
                }
            )
            bold2 = workbook1.add_format(
                {
                    "bold": False,  # 字体加粗
                    "border": 1,  # 单元格边框宽度
                    "align": "center",  # 水平对齐方式
                    "valign": "vcenter",  # 垂直对齐方式
                    "text_wrap": True,  # 是否自动换行
                    "num_format": "0.00",
                }
            )
            bold3 = workbook1.add_format(
                {
                    "bold": False,  # 字体加粗
                    "border": 1,  # 单元格边框宽度
                    "align": "left",  # 水平对齐方式
                    "valign": "vcenter",  # 垂直对齐方式
                    "text_wrap": True,  # 是否自动换行
                }
            )
            # yellow = workbook1.add_format({'fg_color': '#FFEE99'})
            worksheet1.set_column("A:B", 20, cell_format=bold1)  # 设置表格的样式
            worksheet1.set_column("C:C", 45, cell_format=bold1)  # 设置表格的样式
            worksheet1.set_column("D:E", 20, cell_format=bold2)  # 设置表格的样式
            worksheet1.set_column("F:F", 45, cell_format=bold3)  # 设置表格的样式

    def check_if_score_satisfy(self):
        score = self.fund_df.loc[("总分", "总分", "总分"), "分数"]
        if self.template_name != "指数和商品":
            if self.tranche == "优选池":
                if score < 65:
                    self.evaluation_result = "不满足"
                    self.remark.append("评分不满")
            else:
                if score < 75:
                    self.evaluation_result = "不满足"
                    self.remark.append("评分不满")
        else:
            if score < 60:
                self.evaluation_result = "不满足"
                self.remark.append("评分不满")

    def create_word(self):
        document = Document()
        # 增加一级标题, 宋体居中
        Head = document.add_heading("", level=0)  # level设置N级标题
        Head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = Head.add_run(f"{self.fund_name}")
        run.font.name = "楷体"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        ticker_str = (
            ",".join(
                get_related_code_list(self.ticker_symbol, self.include_out_tranche)
            )
            if self.type != "货币"
            else self.ticker_symbol
        )
        self._add_paragraph(document, f"基金代码：{ticker_str}")
        self._add_paragraph(document, f"基金名称：{self.fund_name}")
        self._add_paragraph(document, f"基金经理：{self.manager_name}")
        self._add_paragraph(document, f"评估类型：{self.template_name}")
        self._add_paragraph(document, f"备选池：{self.tranche}")
        self._add_paragraph(
            document,
            f"""评估分数: {self.fund_df.loc[("总分", "总分", "总分"), "分数"]:.2f}""",
        )

        df = self.fund_df.copy().reset_index()
        df2word = DF2WordTable()
        table = df2word.main(document, df)
        self.check_if_score_satisfy()
        # 结论
        if self.if_new == 0:
            if self.evaluation_result == "满足":
                self._add_paragraph(
                    document,
                    f"尽调评估结论：经后续跟踪，基金满足{self.tranche}标准，维持{self.tranche}。",
                )
        else:
            if self.evaluation_result == "满足":
                self._add_paragraph(
                    document, f"基金满足{self.tranche}标准,调入{self.tranche}。"
                )

        self._add_paragraph(document, f"产品标签：{self.type}")
        return document

    def save_word(self):
        document = self.create_word()
        table = document.tables[0]
        # 信息
        cell1 = table.cell(1, 0)
        cell2 = table.cell(3, 0)
        cell1.merge(cell2)
        cell1.text = "信息"

        if self.template_name in ("主动债券", "主动权益"):
            # 定量分析
            self._merge_cells(table, 4, 0, 9, 0)
            # 定性分析
            self._merge_cells(table, 10, 0, 14, 0)

            # 其他
            self._merge_cells(table, 16, 0, 19, 0)
            # 定性分析-基金公司
            self._merge_cells(table, 10, 1, 13, 1)
            self._merge_cells(table, 16, 1, 19, 1)
            # 总分
            self._merge_cells(table, 20, 0, 20, 2)
            # 定性分析-基金经理
            self._merge_cells(table, 6, 1, 9, 1)

        elif self.template_name == "指数和商品":
            # 定量分析
            self._merge_cells(table, 4, 0, 10, 0)
            # 其他
            self._merge_cells(table, 12, 0, 13, 0)
            self._merge_cells(table, 12, 1, 13, 1)
            # 定量分析
            self._merge_cells(table, 5, 1, 6, 1)
            self._merge_cells(table, 7, 1, 8, 1)
            self._merge_cells(table, 9, 1, 10, 1)
            # 总分
            self._merge_cells(table, 14, 0, 14, 2)
        elif self.template_name == "货币":
            # 定量分析
            self._merge_cells(table, 4, 0, 7, 0)
            # 其他
            self._merge_cells(table, 9, 0, 10, 0)
            self._merge_cells(table, 9, 1, 10, 1)
            # 基金产品
            self._merge_cells(table, 5, 1, 7, 1)
            # 总分
            self._merge_cells(table, 11, 0, 11, 2)

        for row in range(len(table.rows)):
            if row == 0:
                for col in range(0, 6):
                    table.cell(row, col).paragraphs[0].runs[0].bold = True
            for col in range(0, 5):
                table.cell(row, col).paragraphs[
                    0
                ].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for col in range(5, 6):
                table.cell(row, col).width = Cm(5)
                table.cell(row, col).paragraphs[
                    0
                ].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._set_table_style(table)
        document.save(f"{self.save_path}.docx")


class EquityFundEvaluation(FundEvaluation):

    def quantivate_part(self):
        self.company_part()
        self.manager_part()
        # 基金规模
        fund_asset_str = f"""基金规模{self.fund_asset:.2f}亿元"""
        fund_asset_score = (
            10
            if self.fund_asset <= 20
            else (0 if self.fund_asset > 120 else 12 - 0.1 * self.fund_asset)
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "产品规模"), "简要评价"] = (
            fund_asset_str
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "产品规模"), "分数"] = (
            fund_asset_score
        )
        # 超额收益
        fund_alpha_str = (
            f"产品年化收益{self.fund_annualised_alpha_dict['基金年化收益']*100:.2f}%,"
            + f"基准年化收益率{self.fund_annualised_alpha_dict['基准年化收益']*100:.2f}%"
            + f",超额收益{self.fund_annualised_alpha_dict['超额']*100:.2f}%"
        )
        fund_alpha_score = 7 + self.fund_annualised_alpha_dict["超额"] / 0.02 * 1
        fund_alpha_score = max(min(15, fund_alpha_score), 0)
        self.fund_df.loc[("定量分析", "基金经理代表产品", "超额收益"), "分数"] = (
            fund_alpha_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "超额收益"), "简要评价"] = (
            fund_alpha_str
        )

        # 最大回撤
        fund_maxdd_str = (
            f"产品最大回撤{self.fund_annualised_alpha_dict['最大回撤']*100:.2f}%,"
            + f"基准年化收益率{self.fund_annualised_alpha_dict['基准最大回撤']*100:.2f}%"
            + f",比值{self.fund_annualised_alpha_dict['回撤比值']:.2f}%"
        )
        fund_maxdd_score = (
            10 + (0.9 - self.fund_annualised_alpha_dict["回撤比值"]) / 0.05 * 1
        )
        fund_maxdd_score = max(min(15, fund_maxdd_score), 0)
        self.fund_df.loc[("定量分析", "基金经理代表产品", "最大回撤"), "分数"] = (
            fund_maxdd_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "最大回撤"), "简要评价"] = (
            fund_maxdd_str
        )

        # 信息比率
        fund_ir_str = f"产品信息比率{self.fund_annualised_alpha_dict['IR']:.2f}%"
        fund_ir_score = (self.fund_annualised_alpha_dict["IR"] - 1) / 0.5 * 1
        fund_ir_score = max(min(5, fund_ir_score), 0)
        self.fund_df.loc[("定量分析", "基金经理代表产品", "信息比率"), "分数"] = (
            fund_ir_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "信息比率"), "简要评价"] = (
            fund_ir_str
        )

    def main(self):
        self.info_part()
        self.quantivate_part()
        self.qualitative_part()
        self.other_part()
        self.fund_df.loc[("总分", "总分", "总分"), "分数"] = self.fund_df["分数"].sum()

        self.save_excel()
        self.save_word()


class BondFundEvaluation(FundEvaluation):

    def quantivate_part(self):
        self.company_part()
        self.manager_part()
        # 基金经理代表产品
        # 基金规模
        fund_asset_str = f"""基金规模{self.fund_asset:.2f}亿元"""
        # 短债基金：产品规模5亿以下为0分，5亿以上每增加5亿得分加1分。
        if self.fund_type_dict["LEVEL_2"] == "短债":
            fund_asset_score = (
                0 if self.fund_asset <= 5 else ((self.fund_asset - 5) / 5 * 1)
            )
        # 除短债基金外：产品规模10亿以下为0分，10亿以上每增加10亿得分加1分
        else:
            fund_asset_score = (
                10 if self.fund_asset <= 10 else ((self.fund_asset - 10) / 10 * 1)
            )
        fund_asset_score = min(fund_asset_score, 10)

        self.fund_df.loc[("定量分析", "基金经理代表产品", "产品规模"), "简要评价"] = (
            fund_asset_str
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "产品规模"), "分数"] = (
            fund_asset_score
        )

        # 绝对收益
        fund_ret_str = f"收益排名{self.fund_perf_rank_dict['绝对业绩']:.2f}%,"
        fund_ret_score = (
            15
            if self.fund_perf_rank_dict["绝对业绩"] <= 20
            else (
                10
                if self.fund_perf_rank_dict["绝对业绩"] <= 50
                else (5 if self.fund_perf_rank_dict["绝对业绩"] <= 70 else 0)
            )
        )
        # fund_alpha_score = max(min(15, fund_alpha_score), 0)
        self.fund_df.loc[("定量分析", "基金经理代表产品", "绝对业绩"), "分数"] = (
            fund_ret_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "绝对业绩"), "简要评价"] = (
            fund_ret_str
        )

        # 最大回撤
        fund_maxdd_str = f"产品最大回撤排名{self.fund_perf_rank_dict['最大回撤']:.2f}%"
        fund_maxdd_score = (
            10
            if self.fund_perf_rank_dict["最大回撤"] <= 20
            else (
                8
                if self.fund_perf_rank_dict["最大回撤"] <= 50
                else (5 if self.fund_perf_rank_dict["最大回撤"] <= 70 else 0)
            )
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "最大回撤"), "分数"] = (
            fund_maxdd_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "最大回撤"), "简要评价"] = (
            fund_maxdd_str
        )

        # 夏普比率
        fund_sr_str = f"产品夏普排名{self.fund_perf_rank_dict['夏普比率']:.2f}%"
        fund_sr_score = (
            10
            if self.fund_perf_rank_dict["夏普比率"] <= 20
            else (
                8
                if self.fund_perf_rank_dict["夏普比率"] <= 50
                else (5 if self.fund_perf_rank_dict["夏普比率"] <= 70 else 0)
            )
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "夏普比率"), "分数"] = (
            fund_sr_score
        )
        self.fund_df.loc[("定量分析", "基金经理代表产品", "夏普比率"), "简要评价"] = (
            fund_sr_str
        )

    def main(self):
        self.info_part()
        self.quantivate_part()
        self.qualitative_part()
        self.other_part()
        self.fund_df.loc[("总分", "总分", "总分"), "分数"] = self.fund_df["分数"].sum()
        self.save_excel()
        self.save_word()


class OtherFundEvaluation(FundEvaluation):
    def company_part(self):
        fund_company_str = (
            f"""{self.management_company}{self.type}基金产品管理规模"""
            + f"""{self.fund_company_type_asset:.2f}亿元,"""
            + f"""在基金公司中排名{self.fund_company_type_rank}"""
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "简要评价"] = (
            fund_company_str
        )

        company_rank_score = (
            20
            if self.fund_company_type_rank <= 20
            else (
                15
                if self.fund_company_type_rank <= 30
                else (
                    10
                    if self.fund_company_type_rank <= 40
                    else (5 if self.fund_company_type_rank <= 50 else 0)
                )
            )
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "分数"] = (
            company_rank_score
        )

    def fund_part(self):
        fund_asset_str = f"""基金规模{self.fund_asset:.2f}亿元"""
        if self.fund_type_dict["LEVEL_2"] in (
            "指数增强",
            "国际(QDII)增强指数型股票基金",
        ):
            # 指数基金
            fund_asset_score = (
                20 if self.fund_asset <= 10 else (20 - (self.fund_asset - 10) / 10 * 1)
            )
            fund_asset_score = max(fund_asset_score, 0)
            self.fund_df.loc[
                ("定量分析", "基金产品-增强指数", "产品规模"), "简要评价"
            ] = fund_asset_str
            self.fund_df.loc[("定量分析", "基金产品-增强指数", "产品规模"), "分数"] = (
                fund_asset_score
            )
            # 超额收益
            fund_alpha_str = (
                f"产品年化收益{self.fund_annualised_alpha_dict['基金年化收益']*100:.2f}%,"
                + f"基准年化收益率{self.fund_annualised_alpha_dict['基准年化收益']*100:.2f}%"
                + f",超额收益{self.fund_annualised_alpha_dict['超额']*100:.2f}%"
            )
            fund_alpha_score = (
                0
                if self.fund_annualised_alpha_dict["超额"] <= 0
                else (self.fund_annualised_alpha_dict["超额"] * 100 / 0.5 * 2)
            )
            fund_alpha_score = min(fund_alpha_score, 30)
            self.fund_df.loc[("定量分析", "基金产品-增强指数", "超额收益"), "分数"] = (
                fund_alpha_score
            )
            self.fund_df.loc[
                ("定量分析", "基金产品-增强指数", "超额收益"), "简要评价"
            ] = fund_alpha_str

        else:
            # 跟踪误差
            tracking_error = (
                get_tracking_error(
                    ticker_symbol=self.ticker_symbol, end_date=self.end_date
                )
                * 100
            )
            tracking_error_str = f"""近1年跟踪误差{tracking_error:.2f}%"""
            if self.fund_type_dict["LEVEL_1"] == "QDII":
                fund_asset_score = (
                    30 if self.fund_asset > 5 else (15 if self.fund_asset > 2 else 0)
                )
                tracking_error_score = (
                    0 if tracking_error > 6 else (6 - tracking_error) / 0.1 * 1
                )

            else:
                fund_asset_score = (
                    30 if self.fund_asset > 40 else (15 if self.fund_asset > 15 else 0)
                )
                tracking_error_score = (
                    0 if tracking_error > 3 else (3 - tracking_error) / 0.1 * 1
                )
            tracking_error_score = min(tracking_error_score, 20)
            if self.fund_type_dict["LEVEL_SUM"] == "商品":
                self.fund_df.loc[
                    ("定量分析", "基金产品-商品基金", "产品规模"), "简要评价"
                ] = fund_asset_str
                self.fund_df.loc[
                    ("定量分析", "基金产品-商品基金", "产品规模"), "分数"
                ] = fund_asset_score
                self.fund_df.loc[
                    ("定量分析", "基金产品-商品基金", "跟踪误差"), "简要评价"
                ] = tracking_error_str
                self.fund_df.loc[
                    ("定量分析", "基金产品-商品基金", "跟踪误差"), "分数"
                ] = tracking_error_score

            else:
                self.fund_df.loc[
                    ("定量分析", "基金产品-被动指数", "产品规模"), "简要评价"
                ] = fund_asset_str
                self.fund_df.loc[
                    ("定量分析", "基金产品-被动指数", "产品规模"), "分数"
                ] = fund_asset_score
                self.fund_df.loc[
                    ("定量分析", "基金产品-被动指数", "跟踪误差"), "简要评价"
                ] = tracking_error_str
                self.fund_df.loc[
                    ("定量分析", "基金产品-被动指数", "跟踪误差"), "分数"
                ] = tracking_error_score

    def qualitative_part(self):
        qualitative_str = (
            "公司近三年合规表现良好，具备较为完善的内控机制"
            if self.fund_company_desc.empty
            else self.fund_company_desc["COMPLIANCE"].values[0]
        )
        self.fund_df.loc[("定性分析", "基金公司", "合规表现情况"), "分数"] = 30
        self.fund_df.loc[("定性分析", "基金公司", "合规表现情况"), "简要评价"] = (
            qualitative_str
        )

    def main(self):
        self.info_part()
        self.company_part()
        self.fund_part()
        self.qualitative_part()
        self.other_part()
        self.fund_df.loc[("总分", "总分", "总分"), "分数"] = self.fund_df["分数"].sum()
        self.save_excel()
        self.save_word()


class CashFundEvaluation(FundEvaluation):
    def company_part(self):
        fund_company_str = (
            f"""{self.management_company}{self.type}基金产品管理规模"""
            + f"""{self.fund_company_type_asset:.2f}亿元,"""
            + f"""在基金公司中排名{self.fund_company_type_rank}"""
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "简要评价"] = (
            fund_company_str
        )

        company_rank_score = (
            5
            if self.fund_company_type_rank <= 20
            else (3 if self.fund_company_type_rank <= 50 else 0)
        )
        self.fund_df.loc[("定量分析", "基金公司", "管理规模"), "分数"] = (
            company_rank_score
        )

    def fund_part(self):
        fund_asset_str = f"""基金规模{self.fund_asset:.2f}亿元"""
        fund_asset_score = (
            15
            if self.fund_asset >= 200
            else (10 if self.fund_asset >= 100 else (5 if self.fund_asset >= 50 else 0))
        )
        self.fund_df.loc[("定量分析", "基金产品", "产品规模"), "简要评价"] = (
            fund_asset_str
        )
        self.fund_df.loc[("定量分析", "基金产品", "产品规模"), "分数"] = (
            fund_asset_score
        )

        fund_2y_rank = get_fund_2y_rank(self.ticker_symbol, self.end_date)
        fund_2y_rank_score = (
            10 if fund_2y_rank <= 30 else (5 if fund_2y_rank <= 50 else 0)
        )
        fund_2y_rank_str = f"近2年收益相对排名{fund_2y_rank:.2f}%"
        self.fund_df.loc[
            ("定量分析", "基金产品", "业绩相对排名(近8个季度)"), "简要评价"
        ] = fund_2y_rank_str
        self.fund_df.loc[
            ("定量分析", "基金产品", "业绩相对排名(近8个季度)"), "分数"
        ] = fund_2y_rank_score
        fund_total_fee = get_total_fee(self.ticker_symbol)
        fund_total_fee_str = f"基金总费用{fund_total_fee:.2f}%"
        self.fund_df.loc[("定量分析", "基金产品", "合计费率"), "简要评价"] = (
            fund_total_fee_str
        )

        fund_total_fee_score = (
            50
            if fund_total_fee <= 0.4
            else (30 if fund_total_fee <= 0.5 else (20 if fund_total_fee <= 0.6 else 0))
        )
        self.fund_df.loc[("定量分析", "基金产品", "合计费率"), "分数"] = (
            fund_total_fee_score
        )

    def qualitative_part(self):
        self.fund_df.loc[("定性分析", "基金公司", "合规表现"), "分数"] = 20
        self.fund_df.loc[("定性分析", "基金公司", "合规表现"), "简要评价"] = (
            "公司近三年合规表现良好，具备较为完善的内控机制"
            if self.fund_company_desc.empty
            else self.fund_company_desc["COMPLIANCE"].values[0]
        )

    def main(self):
        self.info_part()
        self.company_part()
        self.fund_part()
        self.qualitative_part()
        self.other_part()
        self.fund_df.loc[("总分", "总分", "总分"), "分数"] = self.fund_df["分数"].sum()
        self.save_excel()
        self.save_word()


def get_fund_template(ticker_symbol: str, end_date):
    query_sql = f"""
        SELECT
            c.TICKER_SYMBOL,
        CASE
                WHEN d.LEVEL_SUM IN ( '指数', "商品" ) THEN
                '指数和商品' 
                WHEN c.LEVEL_2 = '指数增强' THEN
                '指数和商品' ELSE d.LEVEL_SUM 
            END AS 'template' 
        FROM
            fund_type_own c 
            JOIN fund_type_sum d ON d.LEVEL_1 = c.LEVEL_1 
            AND d.LEVEL_2 = c.LEVEL_2 
        WHERE
            1 = 1 
            and c.TICKER_SYMBOL = '{ticker_symbol}'
            AND c.report_date = (
            SELECT
                max( report_date ) 
            FROM
                fund_type_own 
        WHERE
            report_date < '{end_date}'
        )
    """
    df = DB_CONN_JJTG_DATA.exec_query(query_sql)
    if df.empty:
        raise ValueError("未找到该基金")
    return df["template"].values[0]


def generate_fund_report(
    ticker_symbol: str,
    end_date: str,
    if_new: str,
    tranche: str,
    include_out_tranche: int = 0,
):
    result_dict = {}
    result_dict["ticker_symbol"] = ticker_symbol
    result_dict["end_date"] = end_date
    result_dict["if_new"] = if_new
    result_dict["tranche"] = tranche
    fund_type = get_fund_template(ticker_symbol, end_date)
    result_dict["type"] = fund_type
    try:
        if fund_type == "货币":
            evaluation = CashFundEvaluation(
                ticker_symbol=ticker_symbol,
                end_date=end_date,
                if_new=if_new,
                tranche=tranche,
                include_out_tranche=include_out_tranche,
            )
        if fund_type == "主动债券":
            evaluation = BondFundEvaluation(
                ticker_symbol=ticker_symbol,
                end_date=end_date,
                if_new=if_new,
                tranche=tranche,
                include_out_tranche=include_out_tranche,
            )

        if fund_type == "指数和商品":
            evaluation = OtherFundEvaluation(
                ticker_symbol=ticker_symbol,
                end_date=end_date,
                if_new=if_new,
                tranche=tranche,
                include_out_tranche=include_out_tranche,
            )
        if fund_type == "主动权益":
            evaluation = EquityFundEvaluation(
                ticker_symbol=ticker_symbol,
                end_date=end_date,
                if_new=if_new,
                tranche=tranche,
                include_out_tranche=include_out_tranche,
            )
    except Exception as e:
        print(e)
        return None
    try:
        evaluation.main()
        result_dict["fund_name"] = evaluation.fund_name
        result_dict["关联代码"] = evaluation.related_code
        result_dict["总分"] = evaluation.fund_df.loc[("总分", "总分", "总分"), "分数"]
        result_dict["是否合格"] = evaluation.evaluation_result
        result_dict["备注"] = ";".join(evaluation.remark) if evaluation.remark else ""
        result_dict["基金经理变更"] = check_fund_manager_change(
            ticker_symbol=ticker_symbol, end_date=end_date
        )
    except Exception as e:
        print(e)
    finally:
        df = pd.DataFrame.from_dict(result_dict, orient="index").T
    return df
